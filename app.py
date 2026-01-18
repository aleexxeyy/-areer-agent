import streamlit as st
from langchain_community.chat_models import ChatOllama
from langchain_core.prompts import PromptTemplate
from langchain.callbacks.streaming_stdout import StreamingStdOutCallbackHandler  # <--- Для вывода в консоль
from pypdf import PdfReader
from docx import Document
import io
import time  # <--- Для замеров времени

# --- 1. Настройка страницы (должна быть первой командой) ---
st.set_page_config(
    page_title="AI Career Coach",
    page_icon="💼",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- CSS Хаки для красоты ---
st.markdown("""
    <style>
    .main {
        background-color: #f9f9f9; 
    }
    .stButton>button {
        width: 100%;
        background-color: #FF4B4B;
        color: white;
        height: 3em;
        border-radius: 10px;
        font-weight: bold;
    }
    .stTextArea textarea {
        border-radius: 10px;
    }
    </style>
""", unsafe_allow_html=True)

# --- 2. Сайдбар (Настройки) ---
with st.sidebar:
    st.image("https://cdn-icons-png.flaticon.com/512/4712/4712009.png", width=80)
    st.title("⚙️ Настройки")
    
    st.markdown("### 🧠 Модель")
    model_name = st.text_input("Имя модели Ollama", value="llama3.2")
    
    st.markdown("### 🌡️ Температура")
    temperature = st.slider("Креативность (0 - робот, 1 - творчество)", 0.0, 1.0, 0.3, 0.1)
    
    st.info(f"Используется локальная модель: **{model_name}**")
    st.divider()
    st.caption("Powered by Ollama & LangChain 🦜🔗")

# --- 3. Функции логики ---

def extract_text_from_pdf(pdf_file):
    """Читает PDF"""
    reader = PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text

def create_docx(text):
    """Генерирует DOCX"""
    doc = Document()
    doc.add_heading('Сопроводительное письмо', 0)
    for para in text.split('\n'):
        if para.strip():
            doc.add_paragraph(para)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def get_llm(model, temp):
    """Создает подключение к Ollama с потоковым выводом в консоль"""
    return ChatOllama(
        model=model, 
        temperature=temp,
        callbacks=[StreamingStdOutCallbackHandler()]
    )

def analyze_resume(resume, job, llm):
    template = """
    Ты опытный Tech Recruiter. Проанализируй резюме кандидата для указанной вакансии.
    
    ВАКАНСИЯ:
    {job}
    
    РЕЗЮМЕ:
    {resume}
    
    Сформируй отчет в формате Markdown:
    1. 📊 **Процент соответствия**: (дай число 0-100% и краткое обоснование)
    2. ✅ **Сильные стороны**: (список топ-3)
    3. ⚠️ **Зоны роста**: (что критично отсутствует)
    4. 💡 **Совет**: (один конкретный совет как улучшить резюме для этой вакансии)
    """
    prompt = PromptTemplate.from_template(template)
    chain = prompt | llm
    return chain.invoke({"job": job, "resume": resume}).content

def generate_letter(resume, job, llm):
    template = """
    Напиши убедительное Cover Letter (сопроводительное письмо) на русском языке.
    Тон: Профессиональный, уверенный, но не высокомерный.
    Структура:
    - Приветствие
    - Почему я подхожу (свяжи опыт из резюме с требованиями вакансии)
    - Заключение и призыв к действию.
    
    ВАКАНСИЯ:
    {job}
    
    РЕЗЮМЕ:
    {resume}
    """
    prompt = PromptTemplate.from_template(template)
    chain = prompt | llm
    return chain.invoke({"job": job, "resume": resume}).content

# --- 4. Основной интерфейс ---

st.title("💼 AI Career Coach")
st.markdown("##### Ваш личный помощник для подготовки к собеседованиям")

# Контейнер для входных данных
with st.container():
    col1, col2 = st.columns([1, 1.2], gap="large")

    with col1:
        st.subheader("1️⃣ Загрузите Резюме")
        uploaded_file = st.file_uploader("Выберите PDF файл", type="pdf")

    with col2:
        st.subheader("2️⃣ Описание Вакансии")
        job_description = st.text_area("Вставьте текст вакансии сюда", height=200, placeholder="Требования, стек, обязанности...")

# Кнопка действия
start_analysis = False
if uploaded_file and job_description:
    st.divider()
    _, btn_col, _ = st.columns([1, 2, 1])
    with btn_col:
        start_analysis = st.button("🚀 Запустить Анализ")
elif not uploaded_file or not job_description:
    st.warning("👆 Загрузите резюме и текст вакансии, чтобы начать.")

# --- 5. Обработка и Вывод ---
if start_analysis:
    llm = get_llm(model_name, temperature)
    
    # Используем st.status для отображения шагов
    with st.status("🚀 Запускаю процессы...", expanded=True) as status:
        
        try:
            start_time = time.time()
            
            # --- ЭТАП 1: Чтение ---
            st.write("📂 Читаю PDF файл...")
            print("\n" + "="*30)
            print("[LOG] 1. Чтение PDF...")
            
            resume_text = extract_text_from_pdf(uploaded_file)
            print(f"[LOG]    -> Успешно. Символов: {len(resume_text)}")
            
            # --- ЭТАП 2: Анализ ---
            st.write("🧠 Анализирую резюме (Смотри текст в консоли!)...")
            print("[LOG] 2. Запуск LLM (Анализ)...")
            
            # invoke запустит потоковый вывод в консоль
            analysis_res = analyze_resume(resume_text, job_description, llm)
            print("\n[LOG]    -> Анализ завершен.")
            
            # --- ЭТАП 3: Письмо ---
            st.write("✍️ Пишу сопроводительное письмо...")
            print("[LOG] 3. Запуск LLM (Письмо)...")
            
            letter_res = generate_letter(resume_text, job_description, llm)
            print("\n[LOG]    -> Письмо готово.")
            
            # Завершение
            status.update(label="✅ Все готово!", state="complete", expanded=False)
            total_time = round(time.time() - start_time, 2)
            print(f"[LOG] --- Готово за {total_time} сек ---")
            print("="*30 + "\n")

            st.success(f"Готово! Обработка заняла: {total_time} сек.")
            
            # --- Вкладки с результатами ---
            tab1, tab2 = st.tabs(["📊 Анализ Резюме", "✉️ Сопроводительное Письмо"])
            
            with tab1:
                st.markdown(analysis_res)
                
            with tab2:
                col_text, col_dl = st.columns([3, 1])
                with col_text:
                    st.text_area("Редактируемый черновик:", value=letter_res, height=400)
                with col_dl:
                    st.info("Нравится текст?")
                    docx = create_docx(letter_res)
                    st.download_button(
                        label="💾 Скачать в Word",
                        data=docx,
                        file_name="Cover_Letter.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

        except Exception as e:
            status.update(label="❌ Ошибка!", state="error")
            st.error(f"Упс! Что-то пошло не так: {e}")
            print(f"\n[ERROR] {e}")